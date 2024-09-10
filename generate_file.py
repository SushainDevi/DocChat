import os
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from docx.shared import Inches

def generate_pdf(content, file_name):
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(file_name, pagesize=letter)
    elements = []
    elements.append(Paragraph(content, styles["BodyText"]))
    doc.build(elements)
    return doc

def generate_doc(content, file_name):
    document = Document()
    document.add_heading('Document Title', 0)
    p = document.add_paragraph(content)
    document.save(file_name)
    return document

def generate_file(content, file_format, file_name):
    if file_format == "pdf":
        return generate_pdf(content, file_name)
    elif file_format == "docx":
        return generate_doc(content, file_name)
    else:
        raise ValueError("Unsupported file format")